#!/usr/bin/env python3
"""
PICC企财险知识库 — 条款提取脚本
从 .docx/.doc 文件中提取条款正文，输出纯文本到 corpus/ 目录
"""
import os
import sys
import shutil
from pathlib import Path
from collections import defaultdict

# ====================== 配置 ======================
SOURCE_DIR = Path.home() / "Desktop/龙虾场/baoxian/企财险条款合集"
CORPUS_DIR = Path.home() / "Desktop/龙虾场/picc-property-insurance-kb/corpus"
TAGGED_DIR = Path.home() / "Desktop/龙虾场/picc-property-insurance-kb/tagged"

def extract_docx_text(filepath):
    """从 .docx 提取纯文本"""
    try:
        import docx
        doc = docx.Document(filepath)
        paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)
        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in paragraphs:
                        paragraphs.append(text)
        return '\n\n'.join(paragraphs)
    except Exception as e:
        print(f"  ⚠️ .docx 提取失败: {filepath.name} — {e}")
        return None

def extract_doc_text(filepath):
    """从 .doc 提取纯文本（尝试用 textract 或 antiword）"""
    try:
        # 尝试用 python-docx（新版 .doc 可能兼容）
        import subprocess
        result = subprocess.run(
            ['textutil', '-convert', 'txt', '-stdout', str(filepath)],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except Exception:
        pass

    # 备用：直接读二进制尝试提取可读文本
    try:
        with open(filepath, 'rb') as f:
            content = f.read()
        # 简单提取：取可打印字符段
        import re
        text = content.decode('utf-8', errors='ignore')
        # 保留中英文和标点
        clean = re.sub(r'[^\u4e00-\u9fff\u3000-\u303f\uff00-\uffefa-zA-Z0-9\s.,;:!?()（）《》""''、。，；：！？\-\d%/℃]', '', text)
        if len(clean) > 100:
            return clean
    except Exception:
        pass

    print(f"  ⚠️ .doc 提取失败: {filepath.name}")
    return None

def identify_clause_category(dir_name):
    """从目录名识别险种分类"""
    # 提取编号和名称
    import re
    match = re.match(r'^(\d+)\.?(.*)$', dir_name)
    if match:
        num = match.group(1)
        name = match.group(2).strip().rstrip('.')
        return num, name
    return '99', dir_name

def find_clause_files(product_dir):
    """在险种目录中查找核心条款文件（.docx 优先，.doc 备用）"""
    clause_files = {
        '条款正文': None,
        '费率表': None,
        '要素表': None,
        '可行性报告': None,
        '对比表': None,
        '费率说明': None,
    }

    for root, dirs, files in os.walk(product_dir):
        for f in files:
            if f.startswith('~$') or f == '.DS_Store':
                continue
            fpath = Path(root) / f

            # 跳过"原条款"目录（旧版本）
            if '原条款' in str(fpath):
                continue

            fname_lower = f.lower()
            if fname_lower.endswith('.docx') or fname_lower.endswith('.doc'):
                if '条款' in f and '附加' not in f and '费率' not in f and '对比' not in f and '要素' not in f and '可行性' not in f and '说明' not in f:
                    if clause_files['条款正文'] is None or f.endswith('.docx'):
                        clause_files['条款正文'] = fpath
                elif '费率表' in f or '基础费率' in f:
                    clause_files['费率表'] = fpath
                elif '要素表' in f:
                    clause_files['要素表'] = fpath
                elif '可行性' in f:
                    clause_files['可行性报告'] = fpath
                elif '对比' in f:
                    clause_files['对比表'] = fpath
                elif '费率' in f and '说明' in f:
                    clause_files['费率说明'] = fpath

    return clause_files

def extract_clause_title(text):
    """从文本中提取条款标题"""
    lines = text.strip().split('\n')
    for i, line in enumerate(lines[:10]):
        line = line.strip()
        if '保险' in line and ('条款' in line or '险' in line):
            if len(line) > 10:
                return line
    # fallback: 首行
    if lines:
        return lines[0].strip()
    return "未知条款"

def extract_articles(text):
    """提取条款条目统计"""
    import re
    articles = re.findall(r'第[一二三四五六七八九十百千万\d]+条', text)
    sections = re.findall(r'(总则|保险标的|保险责任|责任免除|保险价值|保险金额|免赔|保险期间|保险人义务|投保人.*义务|赔偿处理|争议处理|其他事项|附则)', text)
    return len(articles), len(sections)

def process_clause_file(filepath, product_name, doc_type):
    """处理单个条款文件"""
    if filepath is None or not filepath.exists():
        return None

    suffix = filepath.suffix.lower()
    if suffix == '.docx':
        text = extract_docx_text(filepath)
    elif suffix == '.doc':
        text = extract_doc_text(filepath)
    else:
        return None

    if text is None:
        return None

    title = extract_clause_title(text)
    article_count, section_count = extract_articles(text)

    return {
        'title': title,
        'text': text,
        'char_count': len(text),
        'article_count': article_count,
        'section_count': section_count,
        'doc_type': doc_type,
    }

def main():
    print("=" * 60)
    print("PICC 企财险知识库 — 条款提取工具")
    print("=" * 60)

    # 清空输出目录
    if CORPUS_DIR.exists():
        shutil.rmtree(CORPUS_DIR)
    CORPUS_DIR.mkdir(parents=True, exist_ok=True)

    # 扫描险种目录
    product_dirs = sorted(SOURCE_DIR.iterdir())
    results = []
    total_chars = 0

    for product_dir in product_dirs:
        if not product_dir.is_dir() or product_dir.name.startswith('.'):
            continue

        num, name = identify_clause_category(product_dir.name)
        print(f"\n📁 [{num}] {name}")

        clause_files = find_clause_files(product_dir)

        if clause_files['条款正文'] is None:
            print(f"  ⚠️ 未找到条款正文文件")
            continue

        # 处理条款正文
        result = process_clause_file(clause_files['条款正文'], name, '条款正文')
        if result is None:
            print(f"  ❌ 条款提取失败")
            continue

        # 保存到 corpus
        safe_name = name.replace('/', '_').replace('\\', '_')
        product_out_dir = CORPUS_DIR / f"{num}_{safe_name}"
        product_out_dir.mkdir(parents=True, exist_ok=True)

        # 保存纯文本
        txt_path = product_out_dir / "条款正文.txt"
        txt_path.write_text(result['text'], encoding='utf-8')

        # 保存元信息
        import json
        meta = {
            'product_id': num,
            'product_name': name,
            'title': result['title'],
            'char_count': result['char_count'],
            'article_count': result['article_count'],
            'section_count': result['section_count'],
            'source_file': str(clause_files['条款正文'].relative_to(SOURCE_DIR)),
        }
        meta_path = product_out_dir / "meta.json"
        meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding='utf-8')

        total_chars += result['char_count']

        print(f"  ✅ {result['title']}")
        print(f"     {result['char_count']} 字符 | {result['article_count']} 条 | {result['section_count']} 个章节")

        results.append(meta)

    # 生成汇总
    print(f"\n{'=' * 60}")
    print(f"📊 提取汇总")
    print(f"   险种数量: {len(results)}")
    print(f"   总字符数: {total_chars:,}")
    print(f"   输出目录: {CORPUS_DIR}")

    # 生成产品索引
    index = {
        'project': 'PICC企财险知识库',
        'total_products': len(results),
        'total_characters': total_chars,
        'products': results,
    }
    index_path = CORPUS_DIR / "index.json"
    index_path.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f"   产品索引: {index_path}")

if __name__ == '__main__':
    main()
